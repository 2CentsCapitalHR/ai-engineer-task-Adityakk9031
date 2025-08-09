import os
import io
import json
import re
import math
import traceback
from datetime import datetime
from typing import List, Dict, Tuple, Optional

import streamlit as st
import numpy as np
from docx import Document
from docx.shared import RGBColor

try:
    from google import genai
    GENAI_CLIENT_AVAILABLE = True
except Exception:
    genai = None
    GENAI_CLIENT_AVAILABLE = False

EMBEDDING_MODEL = "models/embedding-001"
RAG_MODEL = "models/gemini-2.5-flash"
TOP_K = 3
CHUNK_SIZE = 1000

CHECKLISTS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors",
    ]
}

DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "articles"],
    "Memorandum of Association": ["memorandum of association", "memorandum", "moa"],
    "Incorporation Application Form": ["incorporation application", "application for incorporation"],
    "UBO Declaration Form": ["ubo declaration", "ultimate beneficial owner"],
    "Register of Members and Directors": ["register of members", "register of directors", "register of members and directors"],
}

def docx_to_text_and_docobj(docx_bytes: bytes) -> Tuple[str, Document]:
    doc = Document(io.BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n\n".join(paragraphs), doc

def chunk_text(text: str, size: int = CHUNK_SIZE) -> List[str]:
    if not text:
        return []
    paragraphs = text.split("\n\n")
    chunks = []
    cur = []
    cur_len = 0
    for p in paragraphs:
        pl = len(p)
        if cur_len + pl > size and cur:
            chunks.append("\n\n".join(cur))
            cur = [p]
            cur_len = pl
        else:
            cur.append(p)
            cur_len += pl
    if cur:
        chunks.append("\n\n".join(cur))
    return chunks

def l2_norm(vec):
    return math.sqrt(sum(x*x for x in vec))

def cosine_sim(a: List[float], b: List[float]) -> float:
    denom = (l2_norm(a) * l2_norm(b))
    if denom == 0:
        return 0.0
    return float(np.dot(np.array(a), np.array(b)) / denom)

def ensure_genai_ready() -> Tuple[bool, Optional[str]]:
    key = os.environ.get("GEMINI_API_KEY")
    if not key:
        return False, "GEMINI_API_KEY not set in environment."
    if not GENAI_CLIENT_AVAILABLE:
        return False, "google-genai is not installed."
    return True, None

def embed_texts_genai(texts: List[str]) -> List[List[float]]:
    ok, msg = ensure_genai_ready()
    if not ok:
        raise RuntimeError(msg)
    try:
        client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
        resp = client.models.embed_content(model=EMBEDDING_MODEL, contents=texts)
        if not hasattr(resp, "embeddings"):
            if isinstance(resp, dict) and "embeddings" in resp:
                emb_list = resp["embeddings"]
                return [list(e["values"]) if isinstance(e, dict) and "values" in e else list(e) for e in emb_list]
            raise RuntimeError("Unexpected embedding response shape.")
        vecs = []
        for emb in resp.embeddings:
            vals = getattr(emb, "values", None)
            if vals is None and isinstance(emb, dict):
                vals = emb.get("values") or emb.get("embedding") or emb.get("vector")
            if vals is None:
                raise RuntimeError("Embedding item missing values.")
            vecs.append(list(vals))
        return vecs
    except Exception as e:
        raise RuntimeError(f"Embedding call failed: {e}")

def generate_with_genai(prompt: str) -> str:
    ok, msg = ensure_genai_ready()
    if not ok:
        raise RuntimeError(msg)
    try:
        client = genai.Client(api_key=os.environ["GEMINI_API_KEY"])
        resp = client.models.generate_content(model=RAG_MODEL, contents=prompt)
        text = getattr(resp, "text", None)
        if text:
            return text
        if hasattr(resp, "candidates") and resp.candidates:
            candidate = resp.candidates[0]
            try:
                return candidate.content.parts[0].text
            except Exception:
                return str(candidate)
        if isinstance(resp, dict) and "candidates" in resp and resp["candidates"]:
            return json.dumps(resp["candidates"][0])
        return str(resp)
    except Exception as e:
        raise RuntimeError(f"Generation call failed: {e}")

class SimpleRagIndex:
    def __init__(self):
        self.passages: List[Dict] = []

    def add_document(self, title: str, text: str, source: str = ""):
        chunks = chunk_text(text)
        for i, ch in enumerate(chunks):
            self.passages.append({"id": f"{title}__{i}", "text": ch, "source": source, "embedding": None})

    def embed_all(self) -> None:
        texts = [p["text"] for p in self.passages]
        if not texts:
            return
        vecs = embed_texts_genai(texts)
        if len(vecs) != len(self.passages):
            if len(vecs) == 1:
                vecs = [vecs[0]] * len(self.passages)
            else:
                raise RuntimeError("Embedding count mismatch")
        for p, v in zip(self.passages, vecs):
            p["embedding"] = v

    def retrieve(self, query: str, top_k: int = TOP_K) -> List[Dict]:
        if not self.passages or any(p.get("embedding") is None for p in self.passages):
            return []
        qvec = embed_texts_genai([query])[0]
        sims = [(cosine_sim(qvec, p["embedding"]), p) for p in self.passages]
        sims_sorted = sorted(sims, key=lambda x: x[0], reverse=True)
        return [p for score, p in sims_sorted[:top_k]]

def detect_document_types(text: str) -> List[str]:
    found = set()
    t = text.lower()
    for docname, kws in DOC_TYPE_KEYWORDS.items():
        for kw in kws:
            if kw in t:
                found.add(docname)
                break
    return list(found) if found else ["Unknown Document Type"]

def check_for_red_flags(text: str) -> List[Dict]:
    issues = []
    t = text.lower()
    if "adgm" not in t and "abu dhabi global market" not in t:
        issues.append({
            "section": "jurisdiction",
            "issue": "Jurisdiction clause does not specify ADGM",
            "severity": "High",
            "suggestion": "Update jurisdiction to 'Abu Dhabi Global Market (ADGM)'.",
            "citation": None
        })
    if not any(s in t for s in ["signature", "signed", "signature:"]):
        issues.append({
            "section": "signature",
            "issue": "No signature block detected",
            "severity": "Medium",
            "suggestion": "Add a signatory section with name, capacity, and date.",
            "citation": None
        })
    if " may " in t:
        issues.append({
            "section": "language",
            "issue": "Ambiguous use of 'may' detected; could be non-binding",
            "severity": "Low",
            "suggestion": "Consider replacing 'may' with clearer mandatory language if intended.",
            "citation": None
        })
    return issues

def merge_issues(rule: List[Dict], ai: List[Dict]) -> List[Dict]:
    out = []
    seen = set()
    for it in (rule + ai):
        key = (it.get("section",""), (it.get("issue","") or "").strip().lower())
        if key in seen:
            for ex in out:
                ek = (ex.get("section",""), (ex.get("issue","") or "").strip().lower())
                if ek == key:
                    if it.get("citation") and not ex.get("citation"):
                        ex["citation"] = it.get("citation")
                    break
        else:
            seen.add(key)
            out.append(it.copy())
    return out

def add_inline_comment_to_doc(doc: Document, paragraph_match: str, comment_text: str):
    lowered = paragraph_match.lower()
    for para in doc.paragraphs:
        if lowered in (para.text or "").lower():
            run = para.add_run(" ")
            run.add_text("")
            review_run = para.add_run(f"[REVIEW: {comment_text}]")
            review_run.font.color.rgb = RGBColor(0xB2, 0x00, 0x00)
            review_run.bold = True

st.set_page_config(page_title="ADGM Corporate Agent", layout="centered")
st.title("ADGM Corporate Agent — RAG + Rule-based Demo")

if "logs" not in st.session_state:
    st.session_state["logs"] = []

def log(msg: str):
    st.session_state["logs"].append(f"{datetime.utcnow().isoformat()}Z — {msg}")

with st.expander("Debug log (latest at bottom)"):
    for line in st.session_state["logs"][-20:]:
        st.text(line)

with st.expander("1) Upload ADGM reference docs (DOCX / TXT)"):
    ref_files = st.file_uploader("ADGM reference files (optional for RAG)", accept_multiple_files=True, type=["docx","txt"])

with st.expander("2) Upload documents to review (.docx)"):
    user_files = st.file_uploader("User .docx files", accept_multiple_files=True, type=["docx"])

build_btn = st.button("Build/Refresh RAG index (from reference docs)")

if "rag_index" not in st.session_state:
    st.session_state["rag_index"] = SimpleRagIndex()
if "index_built" not in st.session_state:
    st.session_state["index_built"] = False

if build_btn:
    st.session_state["rag_index"] = SimpleRagIndex()
    if not ref_files:
        st.warning("No reference files uploaded.")
        log("Build requested but no reference files.")
    else:
        log(f"Indexing {len(ref_files)} reference files...")
        for rf in ref_files:
            raw = rf.read()
            try:
                text, _ = docx_to_text_and_docobj(raw)
            except Exception:
                text = raw.decode("utf-8", errors="ignore")
            st.session_state["rag_index"].add_document(rf.name, text, source=rf.name)
        try:
            log("Creating embeddings...")
            st.session_state["rag_index"].embed_all()
            st.session_state["index_built"] = True
            st.success("RAG index built.")
        except Exception as e:
            st.session_state["index_built"] = False
            st.error("Failed to create embeddings.")
            log(f"Embedding error: {e}")
            tb = traceback.format_exc().splitlines()[:6]
            for line in tb:
                log(line)

if user_files:
    st.info(f"Analyzing {len(user_files)} file(s)...")
    report = {
        "process": None,
        "documents_uploaded": len(user_files),
        "required_documents": None,
        "missing_documents": [],
        "issues_found": [],
        "timestamp": datetime.utcnow().isoformat() + "Z"
    }
    reviewed_outputs = []
    all_detected = []
    for uf in user_files:
        raw = uf.read()
        text, doc_obj = docx_to_text_and_docobj(raw)
        detected = detect_document_types(text)
        all_detected.append({"filename": uf.name, "detected": detected})
        log(f"Detected for {uf.name}: {detected}")
        rule_issues = check_for_red_flags(text)
        ai_issues = []
        if st.session_state.get("index_built", False):
            try:
                excerpt = text[:1500] if len(text) > 1500 else text
                passages = st.session_state["rag_index"].retrieve(excerpt, top_k=TOP_K)
                if passages:
                    ctx = ""
                    for i, p in enumerate(passages, start=1):
                        ctx += f"[Passage {i} | Source: {p.get('source','')}]\\n{p['text']}\\n\\n"
                    prompt = (
                        "You are an ADGM compliance reviewer. Using ONLY the ADGM passages in CONTEXT, "
                        "analyze the DOCUMENT excerpt and return STRICT JSON with a single key 'issues'.\n\n"
                        f"CONTEXT:\n{ctx}\n\nDOCUMENT_EXCERPT:\n{excerpt}\n\n"
                        "Return only valid JSON."
                    )
                    out = generate_with_genai(prompt)
                    m = re.search(r"\{.*\}$", out, re.S)
                    json_text = m.group(0) if m else out
                    parsed = json.loads(json_text)
                    for gi in parsed.get("issues", []):
                        ai_issues.append({
                            "section": gi.get("section","(from AI)"),
                            "issue": gi.get("issue",""),
                            "severity": gi.get("severity","Medium"),
                            "suggestion": gi.get("suggestion",""),
                            "citation": gi.get("citation")
                        })
            except Exception as e:
                log(f"RAG step failed for {uf.name}: {e}")
                tb = traceback.format_exc().splitlines()[:6]
                for line in tb:
                    log(line)
                st.warning(f"RAG step failed for {uf.name}, rule-based only.")
        merged = merge_issues(rule_issues, ai_issues)
        for it in merged:
            match_phrase = ""
            if "juris" in (it.get("section","") or "").lower():
                if "jurisdiction" in text.lower():
                    match_phrase = "jurisdiction"
                elif "governing law" in text.lower():
                    match_phrase = "governing law"
            if not match_phrase and "signatur" in (it.get("section","") or "").lower():
                if "signature" in text.lower():
                    match_phrase = "signature"
            if not match_phrase and "may" in it.get("issue","").lower():
                match_phrase = "may"
            if not match_phrase:
                match_phrase = (text[:80].split("\n\n")[0])[:40]
            comment_text = f"{it.get('issue')} | Suggestion: {it.get('suggestion')}"
            add_inline_comment_to_doc(doc_obj, match_phrase, comment_text)
        for it in merged:
            it["document"] = uf.name
        report["issues_found"].extend(merged)
        outb = io.BytesIO()
        doc_obj.save(outb)
        outb.seek(0)
        reviewed_outputs.append({"filename": f"reviewed_{uf.name}", "bytes": outb.read()})
    types_present = {t for f in all_detected for t in f['detected']}
    if any(x in types_present for x in CHECKLISTS["Company Incorporation"]):
        report["process"] = "Company Incorporation"
        report["required_documents"] = len(CHECKLISTS["Company Incorporation"])
        required = set(CHECKLISTS["Company Incorporation"])
        present = set([t for f in all_detected for t in f['detected'] if t in required])
        missing = list(required - present)
        report["missing_documents"] = missing
        if missing:
            st.warning(f"It appears you're attempting Company Incorporation. Missing: {', '.join(missing)}")
            log(f"Missing documents: {missing}")
    st.subheader("Summary")
    st.json(report)
    st.download_button("Download JSON report", data=json.dumps(report, indent=2).encode("utf-8"),
                       file_name="adgm_review_report.json", mime="application/json")
    for rd in reviewed_outputs:
        st.download_button(f"Download {rd['filename']}", data=rd['bytes'], file_name=rd['filename'],
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.success("Analysis complete.")
else:
    st.info("Upload user .docx files to analyze.")
